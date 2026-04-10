from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.services.types import ExtractedDocument, ExtractedSegment

SUPPORTED_TYPES: dict[str, str] = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
}


class DocumentParseError(RuntimeError):
    def __init__(self, message: str, code: str = "document_parse_failed") -> None:
        super().__init__(message)
        self.code = code


class DocumentExtractor:
    def extract(self, filename: str, contents: bytes) -> ExtractedDocument:
        extension = Path(filename).suffix.lower()
        if extension == ".txt":
            return self._extract_txt(contents)
        if extension == ".pdf":
            return self._extract_pdf(contents)
        if extension == ".docx":
            return self._extract_docx(contents)
        raise DocumentParseError("Unsupported file type.", code="unsupported_file_type")

    def _extract_txt(self, contents: bytes) -> ExtractedDocument:
        text = contents.decode("utf-8", errors="replace").strip()
        normalized = self._normalize_text(text)
        if not normalized:
            raise DocumentParseError("TXT file did not contain readable text.", code="empty_text")

        segment = ExtractedSegment(
            segment_id="text-body",
            locator="text body",
            text=normalized,
            metadata={"source_type": "txt"},
        )
        return ExtractedDocument(
            parser="txt-basic",
            full_text=normalized,
            segments=[segment],
            metadata={"segment_count": 1},
        )

    def _extract_pdf(self, contents: bytes) -> ExtractedDocument:
        try:
            reader = PdfReader(BytesIO(contents))
        except Exception as exc:  # pragma: no cover - parser failure branch
            raise DocumentParseError(f"PDF could not be opened: {exc}", code="pdf_open_failed") from exc

        segments: list[ExtractedSegment] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = self._normalize_text(page.extract_text() or "")
            if not text:
                continue
            segments.append(
                ExtractedSegment(
                    segment_id=f"page-{page_number}",
                    locator=f"page {page_number}",
                    text=text,
                    metadata={"page_number": page_number, "source_type": "pdf-page"},
                )
            )

        if not segments:
            raise DocumentParseError(
                "PDF opened successfully but no readable text was extracted.",
                code="pdf_no_text",
            )

        return ExtractedDocument(
            parser="pypdf",
            full_text="\n\n".join(segment.text for segment in segments),
            segments=segments,
            metadata={"page_count": len(reader.pages), "segment_count": len(segments)},
        )

    def _extract_docx(self, contents: bytes) -> ExtractedDocument:
        try:
            document = DocxDocument(BytesIO(contents))
        except Exception as exc:  # pragma: no cover - parser failure branch
            raise DocumentParseError(f"DOCX could not be opened: {exc}", code="docx_open_failed") from exc

        segments: list[ExtractedSegment] = []
        for paragraph_number, paragraph in enumerate(document.paragraphs, start=1):
            text = self._normalize_text(paragraph.text)
            if not text:
                continue
            segments.append(
                ExtractedSegment(
                    segment_id=f"paragraph-{paragraph_number}",
                    locator=f"paragraph {paragraph_number}",
                    text=text,
                    metadata={"paragraph_number": paragraph_number, "source_type": "docx-paragraph"},
                )
            )

        for table_number, table in enumerate(document.tables, start=1):
            for row_number, row in enumerate(table.rows, start=1):
                cells = [self._normalize_text(cell.text) for cell in row.cells]
                cells = [cell for cell in cells if cell]
                if not cells:
                    continue
                segments.append(
                    ExtractedSegment(
                        segment_id=f"table-{table_number}-row-{row_number}",
                        locator=f"table {table_number} row {row_number}",
                        text=" | ".join(cells),
                        metadata={
                            "table_number": table_number,
                            "row_number": row_number,
                            "source_type": "docx-table-row",
                        },
                    )
                )

        if not segments:
            raise DocumentParseError(
                "DOCX opened successfully but no readable text was extracted.",
                code="docx_no_text",
            )

        return ExtractedDocument(
            parser="python-docx",
            full_text="\n\n".join(segment.text for segment in segments),
            segments=segments,
            metadata={"segment_count": len(segments)},
        )

    def _normalize_text(self, text: str) -> str:
        return " ".join(text.replace("\x00", " ").split()).strip()
